import os
import json  # ★ v3.1 — Creator JSON 로드용
from datetime import datetime

import streamlit as st
from anthropic import Anthropic

from prompt import (
    SYSTEM_PROMPT,
    GENRE_RULES,
    BEATS_15,
    ACT_BEATS,
    build_scene_plan_prompt,
    build_extract_elements_prompt,
    build_write_beat_prompt,
    build_rewrite_prompt,
    extract_from_creator_json,  # ★ v3.1 신규
    ENGINE_VERSION,              # ★ v3.1
    ENGINE_BUILD_DATE,           # ★ v3.1
)

ANTHROPIC_MODEL_WRITE = "claude-opus-4-6"      # 집필 (비트 쓰기, 다시 쓰기) — 최고 품질
ANTHROPIC_MODEL_PLAN  = "claude-sonnet-4-6"    # 구조 작업 (씬 플랜, 요소 추출) — 비용 효율


# ═══════════════════════════════════════════════════════════
# ★ v3.1.2 — 다운로드 파일명 생성 헬퍼
# 제목을 우선 사용하되, 비어 있으면 장르로 폴백.
# 윈도우/맥 모두에서 문제가 될 수 있는 문자를 안전하게 정제.
# ═══════════════════════════════════════════════════════════

def _sanitize_filename_segment(text: str, max_len: int = 40) -> str:
    """파일명에 쓸 수 있도록 텍스트 정제.
    - 파일시스템 금지 문자 제거 (/ \\ : * ? " < > |)
    - 공백은 언더스코어로 변환
    - 선행/후행 공백·점 제거 (윈도우 금지)
    - 너무 길면 max_len까지 잘라냄
    """
    if not text:
        return ""
    s = str(text).strip()
    # 파일시스템 금지 문자 제거
    for ch in r'/\:*?"<>|':
        s = s.replace(ch, "")
    # 연속 공백 → 단일 언더스코어
    s = "_".join(s.split())
    # 선행·후행 점 제거 (윈도우 이슈)
    s = s.strip(". ")
    if len(s) > max_len:
        s = s[:max_len].rstrip("_")
    return s


def _build_download_filename(title: str, genre: str, ext: str) -> str:
    """다운로드 파일명 생성. 제목 우선, 장르는 폴백."""
    stem = _sanitize_filename_segment(title)
    if not stem:
        stem = _sanitize_filename_segment(genre) or "screenplay"
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    return f"{stem}_{timestamp}.{ext}"


# ═══════════════════════════════════════════════════════════
# ★ v3.1.3 — 지문(액션 라인) 자동 분단 헬퍼
# AI가 6~9문장을 한 문단에 몰아넣은 경우, 의미 비트 단위로 분단.
# AI의 시적 의도가 명확한 경우(짧은 단락)는 절대 건드리지 않음.
# ═══════════════════════════════════════════════════════════

# 분단 임계값 — 데이터 분석 기반 (실제 결과물 481개 단락 분포 검증)
_ACTION_SPLIT_CHAR_THRESHOLD = 150     # 자수 임계
_ACTION_SPLIT_SENTENCE_THRESHOLD = 7   # 문장수 임계
_ACTION_SPLIT_HARD_CHARS = 240         # 하드 임계 (이 이상은 무조건 분단)
_ACTION_SPLIT_HARD_SENTENCES = 9       # 하드 문장수 임계


def _split_sentences(text: str):
    """한국어 지문을 문장 단위로 쪼갠다. 마침표·물음표·느낌표 + 공백 기준.
    문장 부호 뒤 공백이 없는 경우(약어 등)는 분리하지 않는다.
    """
    import re as _re
    # (?<=[.!?])\s+ 로 분리하되, 한국어 마침표 뒤 공백을 기준
    parts = _re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in parts if s.strip()]


def _detect_paragraph_break_index(sentences: list) -> int:
    """
    문장 리스트에서 가장 자연스러운 분단 위치(인덱스)를 찾는다.
    분단 트리거 4종에 따라 우선순위로 탐색.
    찾지 못하면 -1 반환.
    
    트리거 우선순위:
      1) 시간 압축 종료: "수업이 진행된다" / "수업이 끝난다" 같은 큰 사건 변화
      2) 동작 주체 변경: 단체 → 개인, 또는 인물 ↔ 다른 인물
      3) 카메라 시점 변경: 인물 동작 ↔ 정물·공간 묘사 (지시 대명사 없는 명사구 시작)
    """
    import re as _re
    n = len(sentences)
    if n < 4:
        return -1
    
    # 검색 범위: 양 끝 1문장씩 제외 (너무 짧은 분단 방지)
    candidates = []
    
    for i in range(2, n - 1):
        cur = sentences[i]
        prev = sentences[i - 1]
        score = 0
        
        # 트리거 1: 시간 압축 / 큰 상황 전환 키워드
        time_break_patterns = [
            r'수업이?\s*(끝나|진행)',
            r'시간이\s*(흐른|지난|경과)',
            r'(다음\s*날|이튿날|새벽|아침|저녁|밤)',
            r'몇\s*(분|시간|일)\s*(후|뒤)',
            r'(직후|잠시\s*후|곧)',
        ]
        for pat in time_break_patterns:
            if _re.search(pat, cur):
                score += 10
                break
        
        # 트리거 2: 인물명 + 단독/혼자/홀로 (주체 전환 신호)
        if _re.search(r'(혼자|홀로|단독)', cur):
            score += 6
        
        # 트리거 3: 정물·공간 묘사 시작 (인물 주어 없는 명사구)
        # "긴 테이블 위에...", "벽에...", "창밖에..." 같은 패턴
        space_patterns = [
            r'^(긴\s|넓은\s|좁은\s|텅\s|빈\s|새|작은\s|커다란\s)',
            r'^(테이블|벽|창|문|바닥|천장|복도|골목)\s',
            r'^(아일랜드|카운터|책상|의자|침대|소파)\s',
            r'^[가-힣]+\s위에',  # "~ 위에" 시작
        ]
        prev_has_actor = bool(_re.search(r'[가-힣]{2,4}이\s', prev) or _re.search(r'[가-힣]{2,4}가\s', prev))
        cur_is_space = any(_re.search(pat, cur) for pat in space_patterns)
        if prev_has_actor and cur_is_space:
            score += 5
        
        # 트리거 4: 인물 주체 변경 ("A가 ~한다." → "B가 ~한다.")
        prev_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', prev)
        cur_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', cur)
        if prev_actor and cur_actor and prev_actor.group(1) != cur_actor.group(1):
            score += 4
        
        # 위치 보정: 단락의 정중앙 근처가 가장 자연스러운 분단 위치
        center_distance = abs(i - n // 2)
        position_bonus = max(0, 3 - center_distance)
        score += position_bonus
        
        if score >= 5:
            candidates.append((i, score))
    
    if not candidates:
        return -1
    
    # 가장 점수 높은 위치
    candidates.sort(key=lambda x: -x[1])
    return candidates[0][0]


def _split_action_paragraph(text: str) -> list:
    """
    지문 단락이 임계값을 넘으면 의미 비트 단위로 분할.
    임계값 미만이거나 적절한 분단 지점을 못 찾으면 [text] 그대로 반환.
    
    분단 조건 (둘 중 하나 충족):
      - 자수 >= 150자
      - 자수 >= 100자 AND 문장수 >= 7 (단문 리듬 보존을 위한 하한)
    
    Returns:
        분할된 단락 리스트 (1개 또는 그 이상)
    """
    text = text.strip()
    if not text:
        return [text]
    
    char_len = len(text)
    sentences = _split_sentences(text)
    sent_count = len(sentences)
    
    # 분단 조건 — AI 시적 의도(짧은 단문 리듬) 보존을 위해 자수 하한 추가
    triggered_by_length = char_len >= _ACTION_SPLIT_CHAR_THRESHOLD
    triggered_by_sentence = (char_len >= 100 and sent_count >= _ACTION_SPLIT_SENTENCE_THRESHOLD)
    
    if not (triggered_by_length or triggered_by_sentence):
        return [text]
    
    # 분단 시도
    split_idx = _detect_paragraph_break_index(sentences)
    
    # 분단 지점을 못 찾았는데 하드 임계도 안 넘으면 그대로 두기 (보수적)
    if split_idx < 0:
        if char_len < _ACTION_SPLIT_HARD_CHARS and sent_count < _ACTION_SPLIT_HARD_SENTENCES:
            return [text]
        # 하드 임계 초과면 강제로 중간 분할
        split_idx = sent_count // 2
    
    part1 = ' '.join(sentences[:split_idx])
    part2 = ' '.join(sentences[split_idx:])
    
    # 재귀적으로 part2도 검사 (긴 단락이 3분할 필요한 경우)
    result = [part1] + _split_action_paragraph(part2)
    return result


# ═══════════════════════════════════════════════════════════
# ★ v3.1.5 — PROP CONTINUITY 메모 정제
# AI가 비트 끝에 작성한 [소품 상태 / S#N 종료 시점] INTERNAL 메모를
# 최종 시나리오 본문에서 분리·제거.
# - 메모는 AI가 다음 비트 집필 시 참조하도록 prompt가 활용
# - 그러나 출력 시나리오 본문에는 노출되면 안 됨
# - DOCX 빌더가 텍스트를 처리하기 직전에 strip
# ═══════════════════════════════════════════════════════════

import re as _re_prop


def _strip_prop_state_memos(text: str) -> str:
    """
    텍스트에서 [소품 상태 / ...] 메모 블록을 제거.
    
    제거 대상 패턴:
        [소품 상태 / S#N 종료 시점]
        - 노트(세웅): 가방 안 (수업 중 사용 → 회수)
        - 화환: 입구 (시든 채)
        ...
    
    AI가 작성한 메모는 보통:
    1) 빈 줄 + [소품 상태 ... ] 헤더로 시작
    2) 여러 줄의 '- 항목: 위치' 라인
    3) 다음 빈 줄 또는 텍스트 끝까지
    
    Returns:
        메모가 제거된 텍스트.
    """
    if not text:
        return text
    
    # 패턴 1: 코드블록 안에 들어있는 케이스 (```로 감싼 형태)
    # ```\n[소품 상태 ...]\n- ...\n```
    pattern_codeblock = _re_prop.compile(
        r'```[^\n]*\n\[소품\s*상태[^\]]*\][\s\S]*?```',
        _re_prop.MULTILINE
    )
    text = pattern_codeblock.sub('', text)
    
    # 패턴 2: 일반 텍스트 안의 [소품 상태] 블록
    # [소품 상태 ...] 헤더 + - 로 시작하는 라인들
    # 빈 줄(또는 다른 패턴 시작)이 나올 때까지 모두 제거
    pattern_inline = _re_prop.compile(
        r'\n*\[소품\s*상태[^\]]*\]\s*\n'  # 헤더 라인
        r'(?:[\s]*[-•·][^\n]*\n?)+',       # 항목 라인들 (- 또는 • 또는 ·)
        _re_prop.MULTILINE
    )
    text = pattern_inline.sub('\n', text)
    
    # 패턴 3: 코드블록 + INTERNAL 메모 라벨
    # AI가 가끔 "INTERNAL: 소품 상태" 같은 변형도 사용 가능
    pattern_internal = _re_prop.compile(
        r'\n*\[?(?:INTERNAL|작가\s*노트|작가노트|소품\s*추적)[^\]]*\]?\s*\n'
        r'(?:\[소품\s*상태[^\]]*\]\s*\n)?'
        r'(?:[\s]*[-•·][^\n]*\n?)+',
        _re_prop.IGNORECASE | _re_prop.MULTILINE
    )
    text = pattern_internal.sub('\n', text)
    
    # ★ v3.2.0 — GENRE_BOOSTER_CHECK 태그 제거
    # AI가 비트 끝에 작성하는 자가 검증 메모도 본문 노출 금지.
    pattern_booster = _re_prop.compile(
        r'\n*<GENRE_BOOSTER_CHECK>[\s\S]*?</GENRE_BOOSTER_CHECK>\n*',
        _re_prop.IGNORECASE
    )
    text = pattern_booster.sub('\n', text)
    
    # ★ v3.2.0 — HELPER_CHARACTER_CHECK 태그 제거
    pattern_helper = _re_prop.compile(
        r'\n*<HELPER_CHARACTER_CHECK>[\s\S]*?</HELPER_CHARACTER_CHECK>\n*',
        _re_prop.IGNORECASE
    )
    text = pattern_helper.sub('\n', text)
    
    # ★ v3.2.0 — 닫기 태그 없이 떠도는 자가 검증 헤더 (안전망)
    pattern_check_header = _re_prop.compile(
        r'\n*\[★?\s*비트\s*종료[^\]]*GENRE_BOOSTER_CHECK[^\]]*\][\s\S]*?(?=\n\[|\nS#|\n$|\Z)',
        _re_prop.IGNORECASE
    )
    text = pattern_check_header.sub('\n', text)
    
    # 연속된 빈 줄 정리 (3개 이상 → 2개)
    text = _re_prop.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


# ═══════════════════════════════════════════════════════════
# ★ v3.1.4 — INSERT 시스템 (화면 텍스트 표기)
# 카톡·문자·이메일·유튜브·뉴스 등 화면 인서트를 자동 감지하고
# DOCX에서 들여쓰기+이탤릭으로 시각 분리.
# 
# AI가 prompt.py의 INSERT 시스템 모듈에 따라 작성한 표기를 파싱:
#   형식 A: INSERT — [헤더] / 본문 줄들 / [/INSERT]
#   형식 B: [라벨] '본문'  (한 줄)
#   형식 C: ...지문 안에 자연스럽게 따옴표로 인용...  (인라인)
# ═══════════════════════════════════════════════════════════

import re as _re_insert

# INSERT 키워드 — 형식 B 자동 감지용 (라벨 안의 키워드)
_INSERT_LABEL_KEYWORDS = [
    '카톡', '메신저', '라인', '디스코드', '카카오톡',
    '문자', 'SMS', 'MMS',
    '이메일', '메일',
    '유튜브', 'YouTube', 'youtube', 'TV', '뉴스', '방송',
    'SNS', '인스타', '인스타그램', '페이스북', '트위터', 'X', '틱톡', 'DM',
    '검색', '구글', '네이버', '다음',
    '노트', '일기', '메모', '편지', '손글씨', '쪽지',
    '신문', '잡지', '기사', '헤드라인',
    '자막',
    '알림',
    '핸드폰', '핸드폰 화면', '폰', '폰 화면', '화면',
]


def _is_insert_label(text: str) -> bool:
    """[...] 형식 라벨인지 판단 — 형식 B 감지."""
    text = text.strip()
    if not (text.startswith('[') and ']' in text):
        return False
    # [ ... ] 안의 텍스트 추출
    label_match = _re_insert.match(r'^\[([^\]]+)\]', text)
    if not label_match:
        return False
    label_inner = label_match.group(1)
    # 라벨 안에 INSERT 키워드가 있는지 확인
    return any(kw in label_inner for kw in _INSERT_LABEL_KEYWORDS)


def _parse_insert_blocks(text: str) -> list:
    """
    여러 줄 텍스트를 받아 INSERT 블록과 일반 텍스트로 분리.
    
    Returns:
        [{'type': 'action'|'insert_block'|'insert_label', 'data': ...}, ...]
        - action: 일반 지문 텍스트
        - insert_block: 형식 A (헤더 + 본문 줄들 + 닫기)
        - insert_label: 형식 B (라벨 한 줄)
    """
    if not text or not text.strip():
        return []
    
    lines = text.split('\n')
    items = []
    i = 0
    n = len(lines)
    accumulated_action = []
    
    def flush_action():
        if accumulated_action:
            joined = '\n'.join(accumulated_action).strip()
            if joined:
                items.append({'type': 'action', 'data': joined})
            accumulated_action.clear()
    
    while i < n:
        line = lines[i]
        line_stripped = line.strip()
        
        # 형식 A 시작: INSERT — 또는 INSERT - 또는 INSERT:
        if _re_insert.match(r'^INSERT\s*[—\-:]', line_stripped, _re_insert.IGNORECASE):
            flush_action()
            header = line_stripped
            body_lines = []
            i += 1
            # [/INSERT] 또는 빈 줄 + 새 헤딩까지 본문 수집
            while i < n:
                bl = lines[i].strip()
                # 닫기 태그 — 명시적 종료 (라인 자체가 닫기 태그면 소비하고 나감)
                if _re_insert.match(r'^\[/INSERT\]?$', bl, _re_insert.IGNORECASE):
                    i += 1
                    break
                # 빈 줄 처리
                if not bl:
                    # 다음 비-빈 줄 확인
                    j = i + 1
                    while j < n and not lines[j].strip():
                        j += 1
                    if j >= n:
                        # 파일 끝 — 종료
                        i = j
                        break
                    next_line = lines[j].strip()
                    # 다음 줄이 닫기 태그면 그 사이 빈 줄 건너뛰고 닫기 처리
                    if _re_insert.match(r'^\[/INSERT\]?$', next_line, _re_insert.IGNORECASE):
                        i = j + 1
                        break
                    # 다음 줄이 따옴표로 시작하지 않으면 (= 본문이 끝남) 종료
                    if not _re_insert.match(r"^['\"\u2018\u2019\u201C\u201D]", next_line):
                        i = j  # 다음 일반 텍스트 위치로 점프
                        break
                    # 같은 INSERT 블록의 본문이 빈 줄을 끼고 계속됨
                    i += 1
                    continue
                body_lines.append(bl)
                i += 1
            items.append({
                'type': 'insert_block',
                'data': {
                    'header': header,
                    'body': body_lines,
                }
            })
            continue
        
        # 형식 B: [...] 라벨 — 라벨 안에 키워드가 있을 때
        if _is_insert_label(line_stripped):
            flush_action()
            items.append({'type': 'insert_label', 'data': line_stripped})
            i += 1
            continue
        
        # 떠도는 [/INSERT] 단독 라인은 무시 (안전망)
        if _re_insert.match(r'^\[/INSERT\]?$', line_stripped, _re_insert.IGNORECASE):
            i += 1
            continue
        
        # 일반 지문
        accumulated_action.append(line)
        i += 1
    
    flush_action()
    return items


def _parse_insert_label(text: str) -> tuple:
    """
    형식 B 라벨 한 줄을 (label, body)로 분리.
    예: "[핸드폰 / 카톡] '아빠: 임대료 30프로 올린다.'"
        → ("[핸드폰 / 카톡]", "'아빠: 임대료 30프로 올린다.'")
    """
    m = _re_insert.match(r'^(\[[^\]]+\])\s*(.*)$', text.strip())
    if m:
        return m.group(1), m.group(2).strip()
    return text, ""


# ═══════════════════════════════════════════════════════════



# ─────────────────────────────────────
# Page Config
# ─────────────────────────────────────
st.set_page_config(
    page_title="BLUE JEANS · Writer Engine",
    page_icon="👖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────
# ★ v3.1 — Sidebar Engine Info (Creator Engine과 동일 톤)
# ─────────────────────────────────────
with st.sidebar:
    st.markdown(f"""
    <div style="padding:12px;background:#F0F2FF;border-radius:8px;border-left:3px solid #191970;font-family:'Pretendard',sans-serif;">
        <div style="font-size:.72rem;color:#191970;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">ENGINE INFO</div>
        <div style="font-size:1.05rem;font-weight:700;color:#191970;">Writer Engine</div>
        <div style="font-size:1.25rem;font-weight:900;color:#FFCB05;background:#191970;padding:2px 8px;border-radius:4px;display:inline-block;margin-top:4px;">
            {ENGINE_VERSION}
        </div>
        <div style="font-size:.7rem;color:#666;margin-top:8px;">
            Build: {ENGINE_BUILD_DATE}<br>
            Creator Engine v2.3+ 호환
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.caption("버전이 최신인지 확인하세요.")

# ─────────────────────────────────────
# Custom CSS (Creator Engine 동일 톤)
# ─────────────────────────────────────
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #191970; --y: #FFCB05; --bg: #F7F7F5;
    --card: #FFFFFF; --card-border: #E2E2E0; --t: #1A1A2E;
    --g: #2EC484; --dim: #8E8E99; --light-bg: #EEEEF6;
    --serif: 'Paperlogy', 'Noto Serif KR', 'Georgia', serif;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--body); color: var(--t); -webkit-font-smoothing: antialiased;
}
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important; color: var(--t) !important;
}
.stMarkdown, .stText, .stCode { color: var(--t) !important; }
h1,h2,h3,h4,h5,h6 { color: var(--navy) !important; font-family: var(--heading) !important; }
p, span, label, div, li { color: inherit; }
section[data-testid="stSidebar"] { display: none; }

.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-size: 0.92rem !important;
    padding: 0.65rem 0.85rem !important;
}
.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(25,25,112,0.08) !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important; font-size: 0.85rem !important;
}
.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important; color: var(--t) !important;
    border-color: var(--card-border) !important; border-radius: 8px !important;
}
[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important; color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }
.stTextInput label, .stTextArea label, .stSelectbox label {
    color: var(--t) !important; font-weight: 600 !important;
    font-size: 0.82rem !important; margin-bottom: 0.3rem !important;
}

.stButton > button {
    color: var(--t) !important; border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-weight: 700 !important;
    font-size: 0.88rem !important; padding: 0.55rem 1.2rem !important;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(25,25,112,0.08) !important;
}
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important; color: var(--navy) !important;
    border-color: var(--y) !important; font-weight: 800 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.3) !important;
}
.stDownloadButton > button {
    color: var(--navy) !important; border: 1.5px solid var(--y) !important;
    background-color: var(--y) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-weight: 800 !important;
    font-size: 0.88rem !important; padding: 0.55rem 1.2rem !important;
}
.stExpander, details, details summary {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1px solid var(--card-border) !important; border-radius: 8px !important;
}
details[open] > div { background-color: var(--card) !important; }
.stExpander summary, .stExpander summary span { color: var(--t) !important; }
.stAlert { color: var(--t) !important; border-radius: 8px !important; }
[data-testid="stVerticalBlock"], [data-testid="stHorizontalBlock"],
[data-testid="stColumn"] { background-color: transparent !important; }

.header {
    font-size: 0.85rem; font-weight: 700; color: var(--navy);
    letter-spacing: 0.15em; font-family: var(--heading);
}
.brand-title {
    font-size: 2.6rem; font-weight: 900; color: var(--navy);
    font-family: var(--display); letter-spacing: -0.02em;
    position: relative; display: inline-block;
}
.brand-title::after {
    content: ''; position: absolute; bottom: 2px; left: 0;
    width: 100%; height: 4px; background: var(--y); border-radius: 2px;
}
.sub {
    font-size: 0.7rem; color: var(--dim); letter-spacing: 0.15em;
    margin-top: 0.5rem; margin-bottom: 1.5rem;
}
.callout {
    background: var(--light-bg); border-left: 4px solid var(--navy);
    padding: 0.9rem 1.1rem; margin: 0.5rem 0;
    border-radius: 0 8px 8px 0; font-size: 0.88rem; color: var(--t);
}
.cl {
    color: var(--navy); font-weight: 700; font-size: 0.72rem;
    letter-spacing: 0.03em; margin-bottom: 0.3rem; text-transform: uppercase;
}
.section-header {
    background: var(--y); color: var(--navy);
    padding: 0.6rem 1rem; border-radius: 6px;
    font-weight: 800; font-size: 1rem; font-family: var(--heading);
    margin: 1.5rem 0 0.8rem 0;
    display: flex; justify-content: space-between; align-items: center;
}
.section-header .en {
    font-family: var(--display); font-size: 0.75rem;
    font-weight: 700; letter-spacing: 0.05em; opacity: 0.7;
}
.small-meta {
    font-size: 0.78rem; color: var(--dim);
    margin-top: -0.2rem; margin-bottom: 0.5rem;
}
.beat-tag {
    background: var(--navy); color: var(--y);
    display: inline-block; padding: 0.2rem 0.7rem;
    border-radius: 4px; font-size: 0.78rem; font-weight: 800;
    letter-spacing: 0.04em; margin-bottom: 0.4rem;
}
.act-tag {
    background: var(--navy); color: #fff;
    display: inline-block; padding: 0.25rem 0.8rem;
    border-radius: 4px; font-size: 0.82rem; font-weight: 800;
    letter-spacing: 0.06em;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────
# Session State
# ─────────────────────────────────────
FIELDS = ["title", "logline", "intent", "gns", "characters", "opening_strategy",
          "world", "structure", "scene_design", "treatment", "tone",
          "bjnd_data", "ending_payoff", "ending_payoff_type"]  # ★ v3.1 신규 3개

for k, v in {
    "plan_1막": "", "plan_2막": "", "plan_3막": "",
    "story_elements": "",
    "beats_done": {},
    "current_beat": 1,
    "genre": "범죄/스릴러",
    "fmt": "영화 (장편)",
    "fact_based": False,
    "historical": False,
    "historical_type": "팩션",
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

for f in FIELDS:
    if f not in st.session_state:
        st.session_state[f] = ""

# ─────────────────────────────────────
# Helpers
# ─────────────────────────────────────
def get_client():
    api_key = st.secrets.get("ANTHROPIC_API_KEY", os.getenv("ANTHROPIC_API_KEY"))
    return Anthropic(api_key=api_key) if api_key else None

def stream_ai(prompt: str, tokens: int = 16000, model: str = ""):
    """스트리밍 제너레이터. model 미지정 시 WRITE 모델 사용."""
    use_model = model or ANTHROPIC_MODEL_WRITE
    client = get_client()
    if not client:
        yield "❌ ANTHROPIC_API_KEY가 설정되지 않았습니다."
        return
    try:
        with client.messages.stream(
            model=use_model, max_tokens=tokens,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        ) as stream:
            for text in stream.text_stream:
                yield text
    except Exception as e:
        yield f"\n\n❌ 오류: {e}"

def full_plan() -> str:
    """3막 플랜 합침."""
    return "\n\n".join(
        st.session_state.get(f"plan_{a}", "")
        for a in ["1막", "2막", "3막"]
        if st.session_state.get(f"plan_{a}", "")
    )

def plan_ready() -> bool:
    return all(st.session_state.get(f"plan_{a}", "").strip() for a in ["1막", "2막", "3막"])


# ═══════════════════════════════════════════════════════════
# ★ v3.4.0 — 프로젝트 세션 백업 (저장/불러오기)
# 중간에 멈춰도 처음부터 다시 안 해도 되게.
# ═══════════════════════════════════════════════════════════

# 백업 대상 키 — STEP 1 입력 + STEP 1 설정 + STEP 2 결과 + STEP 3 결과
_BACKUP_KEYS = [
    # STEP 1 입력 (FIELDS 14개)
    "title", "logline", "intent", "gns", "characters", "opening_strategy",
    "world", "structure", "scene_design", "treatment", "tone",
    "bjnd_data", "ending_payoff", "ending_payoff_type",
    # STEP 1 설정
    "genre", "fmt", "fact_based", "historical", "historical_type",
    # STEP 2 결과
    "plan_1막", "plan_2막", "plan_3막", "story_elements",
    # STEP 3 결과
    "beats_done", "current_beat",
]


def export_session_backup() -> bytes:
    """현재 세션 상태를 JSON bytes로 직렬화."""
    payload = {
        "_meta": {
            "engine_version": ENGINE_VERSION,
            "build_date": ENGINE_BUILD_DATE,
            "saved_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
            "title": st.session_state.get("title", ""),
            "genre": st.session_state.get("genre", ""),
            "beats_progress": f"{len(st.session_state.get('beats_done', {}))}/15",
        },
        "session": {k: st.session_state.get(k) for k in _BACKUP_KEYS},
    }
    # beats_done의 키는 int인데 JSON은 str로만 직렬화 가능 → 명시 변환
    beats_done = payload["session"].get("beats_done", {})
    if isinstance(beats_done, dict):
        payload["session"]["beats_done"] = {str(k): v for k, v in beats_done.items()}
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def import_session_backup(raw_bytes: bytes) -> dict:
    """JSON bytes를 받아 세션에 복원. 복원된 메타 정보 dict 반환."""
    raw = raw_bytes.decode("utf-8")
    data = json.loads(raw)
    
    session_data = data.get("session", {})
    meta = data.get("_meta", {})
    
    # 세션 상태에 복원
    for k in _BACKUP_KEYS:
        if k in session_data:
            v = session_data[k]
            # beats_done 키를 다시 int로 (JSON에서 str로 저장됨)
            if k == "beats_done" and isinstance(v, dict):
                v = {int(kk): vv for kk, vv in v.items()}
            st.session_state[k] = v
    
    return meta


def make_backup_filename(title: str, beats_count: int) -> str:
    """백업 파일명 생성 — 제목/진행도/시각 포함."""
    base = (title or "Untitled").strip()[:30]
    # 파일명 안전 처리 (윈도우/맥 공통)
    for ch in '<>:"/\\|?*':
        base = base.replace(ch, "_")
    base = base.replace(" ", "_")
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    return f"WriterEngine_{base}_{beats_count}of15_{ts}.json"


# ═══════════════════════════════════════════════════════════
# ★ v3.5.1 — 한국 시나리오 표준 포맷 후처리
# 지문↔대사 사이에 빈 줄을 자동 삽입한다.
# AI가 prompt.py 지시를 따라 빈 줄을 넣어 출력하는 것이 1차 안전망이고,
# 이 함수는 누락 시 보완하는 2차 안전망이다.
# ═══════════════════════════════════════════════════════════

def _normalize_screenplay_blank_lines(text: str) -> str:
    """시나리오 본문에서 지문↔대사 사이 빈 줄을 보정한다.
    
    규칙:
    - 지문 다음 줄이 대사면 사이에 빈 줄 1개
    - 대사 다음 줄이 지문이면 사이에 빈 줄 1개
    - 같은 화자/다른 화자 대사 연속은 빈 줄 없이 유지
    - 씬 헤딩 직전/직후는 기존 빈 줄 처리 유지
    - 이미 빈 줄이 있으면 추가 삽입 안 함 (중복 방지)
    """
    import re
    
    # 라인 분류 함수
    heading_pat = re.compile(r'^S#\d+', re.UNICODE)
    # 대사 패턴: "캐릭터명\t\t대사" (탭 1~3개 허용)
    dialogue_pat = re.compile(r'^[^\t]+\t{1,}\S', re.UNICODE)
    
    def line_type(line: str) -> str:
        s = line.rstrip()
        if not s.strip():
            return "blank"
        if heading_pat.match(s.strip()):
            return "scene"
        if dialogue_pat.match(s):
            return "dialogue"
        return "action"
    
    lines = text.split('\n')
    result = []
    for idx, line in enumerate(lines):
        cur_type = line_type(line)
        # 이전 의미 있는 라인 타입 찾기 (빈 줄 건너뛰기)
        prev_meaningful = None
        for r_line in reversed(result):
            r_type = line_type(r_line)
            if r_type != "blank":
                prev_meaningful = r_type
                break
        
        # 직전 줄이 빈 줄인지 (이미 분리된 상태인지)
        already_separated = bool(result) and not result[-1].strip()
        
        # 빈 줄 삽입 결정
        need_blank = False
        if prev_meaningful and not already_separated:
            # 지문 → 대사 또는 대사 → 지문 전환
            if (prev_meaningful == "action" and cur_type == "dialogue") or \
               (prev_meaningful == "dialogue" and cur_type == "action"):
                need_blank = True
        
        if need_blank:
            result.append("")
        result.append(line)
    
    return '\n'.join(result)


def make_docx_bytes(genre: str, beats_done: dict, title: str = "",
                    fact_based: bool = False,
                    historical: bool = False,
                    historical_type: str = "") -> bytes:
    """시나리오 DOCX — 한국 표준 시나리오 서식."""
    import re
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from io import BytesIO

    doc = DocxDocument()

    # ── 페이지 설정 (A4, 20mm 마진) ──
    from docx.shared import Mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # ── 기본 스타일: 함초롬바탕 10pt ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "함초롬바탕"
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.space_before = Pt(0)
    # 한글 폰트 설정
    rpr = style_normal.element.rPr
    if rpr is None:
        rpr = style_normal.element.makeelement(qn('w:rPr'), {})
        style_normal.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '함초롬바탕')

    # ── 커스텀 Word 스타일 생성 (나중에 Word 스타일 패널에서 일괄 수정 가능) ──
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm

    def _set_eastasia_font(style_or_run_element, font_name='함초롬바탕'):
        """rPr에 eastAsia 폰트 설정."""
        rpr_elem = style_or_run_element
        rf = rpr_elem.find(qn('w:rFonts'))
        if rf is None:
            rf = rpr_elem.makeelement(qn('w:rFonts'), {})
            rpr_elem.append(rf)
        rf.set(qn('w:eastAsia'), font_name)

    # [스타일 1] 씬번호 — Bold, 위 24pt / 아래 6pt, 1.5줄 간격
    style_scene = doc.styles.add_style('씬번호', WD_STYLE_TYPE.PARAGRAPH)
    style_scene.base_style = doc.styles['Normal']
    style_scene.font.name = '함초롬바탕'
    style_scene.font.size = Pt(11)
    style_scene.font.bold = True
    style_scene.paragraph_format.space_before = Pt(24)
    style_scene.paragraph_format.space_after = Pt(6)
    style_scene.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_scene.element.get_or_add_rPr())

    # [스타일 2] 대사 — Bold, 왼쪽 들여쓰기 1.25cm, 위 8pt / 아래 2pt
    style_dialogue = doc.styles.add_style('대사', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue.base_style = doc.styles['Normal']
    style_dialogue.font.name = '함초롬바탕'
    style_dialogue.font.size = Pt(10)
    style_dialogue.font.bold = True
    style_dialogue.paragraph_format.left_indent = Cm(1.25)
    style_dialogue.paragraph_format.space_before = Pt(8)
    style_dialogue.paragraph_format.space_after = Pt(2)
    style_dialogue.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_dialogue.element.get_or_add_rPr())

    # [스타일 3] 대사연속 — 대사 이어쓰기 (캐릭터명 없이)
    style_dialogue_cont = doc.styles.add_style('대사연속', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue_cont.base_style = style_dialogue
    style_dialogue_cont.paragraph_format.space_before = Pt(0)
    style_dialogue_cont.paragraph_format.space_after = Pt(0)

    # [스타일 4] 지문 — 일반 텍스트, 들여쓰기 없음
    style_action = doc.styles.add_style('지문', WD_STYLE_TYPE.PARAGRAPH)
    style_action.base_style = doc.styles['Normal']
    style_action.font.name = '함초롬바탕'
    style_action.font.size = Pt(10)
    style_action.font.bold = False
    style_action.paragraph_format.space_before = Pt(2)
    style_action.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_action.element.get_or_add_rPr())

    # ★ v3.1.4 — INSERT 전용 스타일 3종
    # [스타일 5] INSERT 헤더 — 작은 대문자 느낌, 중간 들여쓰기, 굵게
    style_insert_header = doc.styles.add_style('인서트헤더', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_header.base_style = doc.styles['Normal']
    style_insert_header.font.name = '함초롬바탕'
    style_insert_header.font.size = Pt(9)
    style_insert_header.font.bold = True
    style_insert_header.paragraph_format.left_indent = Cm(2.55)
    style_insert_header.paragraph_format.space_before = Pt(8)
    style_insert_header.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_insert_header.element.get_or_add_rPr())

    # [스타일 6] INSERT 본문 — 깊은 들여쓰기, 이탤릭, 본문 표시
    style_insert_body = doc.styles.add_style('인서트본문', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_body.base_style = doc.styles['Normal']
    style_insert_body.font.name = '함초롬바탕'
    style_insert_body.font.size = Pt(10)
    style_insert_body.font.italic = True
    style_insert_body.paragraph_format.left_indent = Cm(2.55)
    style_insert_body.paragraph_format.space_before = Pt(2)
    style_insert_body.paragraph_format.space_after = Pt(2)
    style_insert_body.paragraph_format.line_spacing = 1.4
    _set_eastasia_font(style_insert_body.element.get_or_add_rPr())

    # [스타일 7] INSERT 라벨식 — 한 줄 짜리 [라벨] '본문' 형식
    style_insert_label = doc.styles.add_style('인서트라벨', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_label.base_style = doc.styles['Normal']
    style_insert_label.font.name = '함초롬바탕'
    style_insert_label.font.size = Pt(10)
    style_insert_label.paragraph_format.left_indent = Cm(1.42)
    style_insert_label.paragraph_format.space_before = Pt(4)
    style_insert_label.paragraph_format.space_after = Pt(4)
    _set_eastasia_font(style_insert_label.element.get_or_add_rPr())

    # ── 헬퍼 함수 (스타일 기반) ──
    def add_text(text, bold=False, size=None, color=None, align=None):
        """커버 페이지용 범용 텍스트."""
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        if bold:
            r.bold = True
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color
        return p

    def add_scene_heading(text):
        """씬 헤딩 — '씬번호' 스타일 적용."""
        p = doc.add_paragraph(style='씬번호')
        r = p.add_run(text)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_dialogue(char_name, parenthetical, line, continuation=False):
        """대사 — '대사'/'대사연속' 스타일 적용.
        continuation=True이면 캐릭터명 생략.
        
        ★ v3.1.3 — 대사 본문 내 괄호 지시문(예: "(앞치마를 내리며)")은
        run을 분할하여 bold를 해제. 화자 표기의 괄호(예: "(V.O.)")는 유지.
        """
        if continuation:
            p = doc.add_paragraph(style='대사연속')
            speaker_part = "\t\t"
        else:
            p = doc.add_paragraph(style='대사')
            speaker_part = f"{char_name}\t\t"

        # 대사 본문 영역 조립: parenthetical 인자 + 실제 대사
        body_parts = []  # [(text, is_paren), ...] — is_paren=True면 bold 해제
        if parenthetical:
            body_parts.append((f"({parenthetical}) ", True))
        if line:
            # line 자체에도 인라인 괄호가 있을 수 있음 (예: "알겠어. (잠깐) 그치만")
            # 정규식으로 (...) 부분과 그 외 부분을 분리
            import re as _re
            chunks = _re.split(r'(\([^()]*\))', line)
            for chunk in chunks:
                if not chunk:
                    continue
                if chunk.startswith('(') and chunk.endswith(')'):
                    body_parts.append((chunk, True))
                else:
                    body_parts.append((chunk, False))

        # run 1: 화자 영역 (화자명 + 탭 또는 탭만) — 스타일의 bold 그대로 유지
        r_speaker = p.add_run(speaker_part)
        r_speaker.font.name = "함초롬바탕"
        _set_eastasia_font(r_speaker._element.get_or_add_rPr())

        # run 2~N: 대사 본문 — 괄호 부분은 bold=False 명시
        for text, is_paren in body_parts:
            r = p.add_run(text)
            r.font.name = "함초롬바탕"
            _set_eastasia_font(r._element.get_or_add_rPr())
            if is_paren:
                r.bold = False  # 괄호 지시문은 bold 해제

        return p

    def add_blank_line():
        """★ v3.5.1 — 지문↔대사 사이 빈 줄 (한국 시나리오 표준 포맷).
        함초롬바탕 10pt로 통일된 빈 단락. 워드/한글 모두 일관된 높이."""
        p = doc.add_paragraph(style='지문')
        # 빈 run으로 폰트만 지정 (텍스트 없음)
        r = p.add_run("")
        r.font.name = "함초롬바탕"
        r.font.size = Pt(10)
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_action(text):
        """지문 — '지문' 스타일 적용.
        
        ★ v3.1.3 — 긴 단락은 의미 비트 단위로 자동 분단.
        ★ v3.1.4 — INSERT 블록(형식 A·B) 자동 감지 → 전용 스타일로 분기.
        
        AI 시적 의도 보존을 위해 짧은 단락(150자 미만, 7문장 미만)은 그대로 둠.
        """
        # ★ v3.1.4: INSERT 블록 우선 분리
        items = _parse_insert_blocks(text)
        
        first_p = None
        for item in items:
            if item['type'] == 'insert_block':
                # 형식 A — 헤더 + 본문 줄들
                p = add_insert_block(item['data']['header'], item['data']['body'])
            elif item['type'] == 'insert_label':
                # 형식 B — 라벨 한 줄
                p = add_insert_label(item['data'])
            else:
                # 일반 지문 — 분단 알고리즘 적용
                sub_paragraphs = _split_action_paragraph(item['data'])
                p = None
                for sub in sub_paragraphs:
                    sp = doc.add_paragraph(style='지문')
                    r = sp.add_run(sub)
                    r.font.name = "함초롬바탕"
                    _set_eastasia_font(r._element.get_or_add_rPr())
                    if p is None:
                        p = sp
            
            if first_p is None:
                first_p = p
        
        return first_p

    def add_insert_block(header: str, body_lines: list):
        """★ v3.1.4 — INSERT 블록 (형식 A) 렌더링.
        
        헤더(작게·굵게·들여쓰기) + 본문(이탤릭·깊은 들여쓰기) + 자동 빈 줄.
        예시:
            INSERT — 핸드폰 카톡 화면
            
              '아빠: 임대료 30프로 올린다.'
              '아빠: 6개월 안에 결정.'
            
            [/INSERT]
        """
        # 빈 줄 한 개 (시각 분리)
        doc.add_paragraph("")
        
        # 헤더 단락
        first_p = doc.add_paragraph(style='인서트헤더')
        r = first_p.add_run(header.strip())
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        
        # 본문 줄들 (각 줄을 별도 단락으로)
        for line in body_lines:
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='인서트본문')
            r = p.add_run(line)
            r.font.name = "함초롬바탕"
            r.italic = True  # 명시적 이탤릭 (스타일 상속에 더해 안전망)
            _set_eastasia_font(r._element.get_or_add_rPr())
        
        # 닫기 표시 (감독·배우용 시각 종료 신호)
        close_p = doc.add_paragraph(style='인서트헤더')
        cr = close_p.add_run('[/INSERT]')
        cr.font.name = "함초롬바탕"
        _set_eastasia_font(cr._element.get_or_add_rPr())
        
        # 빈 줄 한 개 (시각 분리)
        doc.add_paragraph("")
        
        return first_p

    def add_insert_label(text: str):
        """★ v3.1.4 — 형식 B 라벨식 INSERT 한 줄 렌더링.
        
        예: [핸드폰 / 카톡] '아빠: 임대료 30프로 올린다.'
            ↑ 라벨 부분 (작게)    ↑ 본문 부분 (이탤릭)
        """
        label, body = _parse_insert_label(text)
        
        p = doc.add_paragraph(style='인서트라벨')
        
        # 라벨 부분 (작게, 굵게)
        r_label = p.add_run(label + ' ')
        r_label.font.name = "함초롬바탕"
        r_label.font.size = Pt(9)
        r_label.bold = True
        _set_eastasia_font(r_label._element.get_or_add_rPr())
        
        # 본문 부분 (이탤릭)
        if body:
            r_body = p.add_run(body)
            r_body.font.name = "함초롬바탕"
            r_body.italic = True
            _set_eastasia_font(r_body._element.get_or_add_rPr())
        
        return p

    # ── 커버 페이지 ──
    for _ in range(6):
        doc.add_paragraph("")
    add_text("시나리오", size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    proj_title = title or f"<{genre}>"
    add_text(proj_title, bold=True, size=Pt(24), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_text("기획/제작 | 블루진픽처스", size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x8E, 0x8E, 0x99))
    add_text(f"Writer Engine {ENGINE_VERSION}  ·  {len(beats_done)}/15 비트",
             size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x8E, 0x8E, 0x99))
    doc.add_page_break()

    # ── 면책 자막 페이지 ──
    # 적용 조건:
    #   (1) 실화 배경 기반 작품 (fact_based)
    #   (2) 팩션·퓨전 역사영화 (정통은 실존 사건·인물 중심이라 면책 문구가 맞지 않음)
    _need_disclaimer = fact_based or (
        historical and ("팩션" in (historical_type or "") or "퓨전" in (historical_type or "")
                        or "faction" in (historical_type or "").lower()
                        or "fusion" in (historical_type or "").lower())
    )
    if _need_disclaimer:
        for _ in range(10):
            doc.add_paragraph("")
        add_text("본 작품에 등장하는 인물, 단체, 지명, 상호, 사건은",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("모두 허구이며, 실존하는 것과 관련이 있더라도",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("극적 구성을 위해 각색되었습니다.",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph("")
        add_text("All characters, organizations, places, and events in this work",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        add_text("are fictional. Any resemblance to actual persons or events is",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        add_text("dramatized for narrative purposes.",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        doc.add_page_break()

    # ── 본문 파싱 + 변환 ──
    # 씬 헤딩 패턴: S#숫자. INT./EXT. 또는 그냥 INT./EXT.
    heading_re = re.compile(r'^S?#?\d*\.?\s*(INT\.|EXT\.|INT\./EXT\.)\s*(.+)', re.IGNORECASE)
    # 캐릭터명 패턴: 2칸+ 들여쓰기 + 이름 (1~15자) + 선택적 (V.O.) / (O.S.) / (CONT'D)
    char_re = re.compile(
        r'^\s{2,}([가-힣a-zA-Z\s]{1,15}?)\s*'
        r'(?:\((V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.)\))?\s*$',
        re.IGNORECASE
    )
    # 인라인 대사 패턴: "캐릭터명\t\t(지시)대사" 또는 "캐릭터명 (V.O.)\t\t대사"
    inline_dialogue_re = re.compile(
        r'^([가-힣a-zA-Z\s]{1,15}?)\s*'
        r'(?:\((V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.)\))?\s*'
        r'\t{1,}\s*(?:\(([^)]*)\)\s*)?(.+)',
        re.IGNORECASE
    )
    # 괄호 지시
    paren_re = re.compile(r'^\s{2,}\((.+?)\)\s*$')
    # 구분선/내부메모 (출력 제외)
    divider_re = re.compile(r'^(═{3,}|─{3,}|---)')

    # ═══════════════════════════════════════════════════════════
    # 메타데이터 유출 차단 — 강화된 필터 (v3.4)
    # Writer Engine 프롬프트가 요구하는 "내부 메모" 항목이 본문에 유출되는 버그 차단.
    # 아래 키워드 중 하나라도 매칭되면 해당 라인은 DOCX 본문에서 제외.
    # ═══════════════════════════════════════════════════════════
    META_PREFIX_PATTERNS = [
        # 장르 장치 메타 (10개 장치 x 9장르)
        "premise_engine", "comic_contradiction", "character_comic_flaw",
        "comic_escalation", "line_surprise", "status_comedy",
        "timing_precision", "callback_payoff", "scene_comic_engine",
        "joke_density",
        "fear_anticipation", "uncertainty", "sensory_unease",
        "threat_design", "dread_pacing", "violation_of_safety",
        "image_residue", "vulnerability", "false_relief",
        "terror_escalation",
        "information_asymmetry", "escalation", "clock_device",
        "suspense_peak", "plot_twist", "investigator_obstacle",
        "villain_intelligence", "moral_ambiguity", "red_herring",
        "irreversible_stakes",
        "action_spark", "physical_choreography", "setpiece_scale",
        "hero_signature", "obstacle_escalation", "stakes_personal",
        "counter_attack", "low_point", "final_confrontation",
        "kinetic_rhythm",
        "longing_distance", "touch_hesitation", "romantic_specificity",
        "emotional_subtext", "miscommunication", "emotional_reversal",
        "vulnerability_moment", "physical_chemistry", "obstacle_internal",
        "payoff_emotional",
        "world_rule", "tech_showcase", "awe_moment", "info_drip",
        "human_anchor", "rule_consequence", "visual_wonder",
        "scale_shift", "philosophical_stakes", "discovery_rhythm",
        "magic_rule", "mythic_echo", "threshold_crossing",
        "wonder_image", "sacrifice_price", "prophecy_twist",
        # 비트 메타 공통
        "writer_notes", "plant_payoff_tag", "plant_payoff",
        "scene_meta", "quality_check",
        # 한글 메타 헤더
        "맥거핀", "캐릭터 비밀", "핵심 장소", "모티프", "모티프:",
        "Plant:", "Plant/Payoff", "Payoff:", "Payoff :",
        "서브플롯", "관객 심리", "열린 질문", "Dramatic Irony",
        "Zeigarnik", "보이스 점검", "보이스점검",
        "비트 요약", "비트요약", "비트 구조 유형", "액션 아이디어",
        "서사동력", "작동한 장르", "작동 장르", "핵심 요소",
        "장르 드라이브", "캐릭터 전술", "캐릭터 아크",
        "서브플롯 진행", "강회장 B-Story", "B-Story",
        "민준 서브플롯", "민준 아크",
        # 장르 드라이브 5점 체크 항목 (① ② ③ ④ ⑤ 뒤에 오는 메타 키워드)
        "정보 비대칭", "정보비대칭", "에스컬레이션",
        "적대자", "타이머", "장르 쾌감", "장르쾌감",
        # 추가 메타 헤더 (v3.4.1 보강)
        "캐릭터 전술", "캐릭터전술", "Payoff 회수", "Payoff회수",
        "Plant 유지", "Plant유지", "비밀",
    ]
    # 줄의 시작 부분에 대한 매칭 (불릿/기호 뒤 텍스트)
    META_LINE_RE = re.compile(
        r'^(?:[\s•·\-*⭐★─═]+\**)*(?:[①②③④⑤⑥⑦⑧⑨⑩]\s*)?'
        r'(' + '|'.join(re.escape(p) for p in META_PREFIX_PATTERNS) + r')'
        r'(?:\s|[:\-—.(]|$)',
        re.IGNORECASE
    )

    # (Beat N plant → S#NN payoff) / (S#NN → S#NN) / (전체 plant → S#NN payoff) 같은 개발자 표기
    META_DEV_NOTATION_RE = re.compile(
        r'\((?:Beat\s*\d+|S#\s*\d+|전체|전반|후반)[^)]*(?:plant|payoff|→|->)[^)]*\)',
        re.IGNORECASE
    )
    # "- 설명(S#NN) — 미공개/미등장/공개/열린 채" 같은 단독 dev 코멘트
    META_DEV_COMMENT_RE = re.compile(
        r'^[\s•·\-*]+.*?\(S#\s*\d+(?:[/,]\s*\d+)*\)\s*(?:—|-|–)\s*(?:미공개|미등장|미해결|공개|폭로|열린|유지|부재)',
        re.IGNORECASE
    )
    # "(관객 O, 유진 X) — 미공개/유지" 같은 Dramatic Irony 메타 주석
    META_IRONY_COMMENT_RE = re.compile(
        r'\(관객\s*[OX].*?(?:유진|주인공|캐릭터)\s*[OX][^)]*\)\s*(?:—|-|–)\s*(?:미공개|미등장|미해결|공개|유지)',
    )
    # "· 캐릭터명:" 또는 "- 캐릭터명:" 으로 시작하는 서브플롯 요약
    # (본문 대사와 구분: 대사는 "캐릭터\t\t대사", 요약은 불릿+콜론)
    META_CHARACTER_SUMMARY_RE = re.compile(
        r'^[\s•·\-*]+\s*([가-힣]+(?:[·∙・]\s*[가-힣]+)*(?:\s*커플)?)\s*:',
    )
    # "- 설명 (Beat N plant → S#NN payoff): 내용" — 불릿 + 설명 + 괄호 dev notation + 콜론
    META_BULLET_DEV_RE = re.compile(
        r'^[\s•·\-*]+.*?\((?:Beat\s*\d+|S#\s*\d+)[^)]*\)\s*:',
        re.IGNORECASE
    )

    def is_meta_line(s: str) -> bool:
        """메타데이터 라인 여부 판정."""
        if not s:
            return False
        # 1차: 메타 키워드 직접 매칭
        if META_LINE_RE.match(s):
            return True
        # 2차: snake_case 장르 장치명 (· premise_engine ...)
        m = re.match(r'^[\s•·\-*⭐★─═]+\s*([a-z]+_[a-z_]+)', s)
        if m:
            return True
        # 3차: 개발자 표기 괄호 + 콜론 (- 간판 케이블타이(Beat 1 plant → S#95 payoff):)
        if META_BULLET_DEV_RE.match(s):
            return True
        # 4차: 캐릭터명 + 콜론 서브플롯 요약 (· 오현수·박지영 커플: ...)
        # 주의: 본문 대사는 "캐릭터\t\t대사"이고 불릿 없음 → 오차단 안 됨
        if META_CHARACTER_SUMMARY_RE.match(s):
            return True
        # 5차: 줄 전체에 개발자 표기가 있으면 메타 (예: "... (Beat 3 plant → S#45 payoff) ...")
        if META_DEV_NOTATION_RE.search(s):
            return True
        # 6차: 개발자 코멘트 라인 (예: "- 진호 도면 메모(S#91) — 미공개.")
        if META_DEV_COMMENT_RE.match(s):
            return True
        # 7차: Dramatic Irony 메타 주석 (예: "- 건물 강회장 소유(관객 O, 유진 X) — 미공개.")
        if META_IRONY_COMMENT_RE.search(s):
            return True
        return False

    current_act = ""
    for b_no in sorted(beats_done.keys()):
        b_info = BEATS_15[b_no - 1]

        # ACT 전환
        if b_info["act"] != current_act:
            if current_act:
                doc.add_page_break()
            current_act = b_info["act"]

        text = beats_done[b_no]

        # ★ v3.1.5 — PROP CONTINUITY 메모 자동 제거
        # AI가 비트 끝에 작성한 [소품 상태 / S#N 종료 시점] INTERNAL 메모는
        # 다음 비트 집필용 참조 자료로만 쓰이고, 최종 시나리오 본문에는 노출 안 됨.
        text = _strip_prop_state_memos(text)

        # ═══════════════════════════════════════════════════════════
        # 대사 형식 붕괴 자동 복구 (v3.4 신규)
        # 버그: 긴 컨텍스트에서 AI가 대사 포맷 규칙을 잊고
        #       "캐릭터\n\n대사" 형식으로 출력 (S#89~ 에서 발견됨)
        # 복구: "캐릭터" 단독 라인 + 빈 라인 + 대사 라인 → "캐릭터\t\t대사"
        # ═══════════════════════════════════════════════════════════
        _CHAR_NAMES = {
            '유진', '진호', '세웅', '다은', '강회장', '민준', '박지영', '오현수',
            '이진호', '반세웅', '김사장', '비서', '편집자', '기사', '배달 기사',
            '사장', '민준 엄마', '박씨', '엄마', '아빠', '형', '누나', '아들', '딸',
        }
        _broken_lines = text.split("\n")
        _fixed_lines = []
        _j = 0
        while _j < len(_broken_lines):
            _cur = _broken_lines[_j].strip()
            # 패턴 A: "캐릭터명" 단독 + 빈줄 + 대사 → "캐릭터\t\t대사"
            if (_cur in _CHAR_NAMES and
                _j + 2 < len(_broken_lines) and
                _broken_lines[_j + 1].strip() == "" and
                _broken_lines[_j + 2].strip() and
                not _broken_lines[_j + 2].strip().startswith("S#") and
                _broken_lines[_j + 2].strip() not in _CHAR_NAMES):
                _next_content = _broken_lines[_j + 2].strip()
                # 괄호 지시(예: "(잠깐 생각하고)")가 있으면 다음 줄이 진짜 대사
                if _next_content.startswith("(") and _next_content.endswith(")") and \
                   _j + 4 < len(_broken_lines) and _broken_lines[_j + 3].strip() == "" and \
                   _broken_lines[_j + 4].strip():
                    _fixed_lines.append(f"{_cur}\t\t{_next_content} {_broken_lines[_j + 4].strip()}")
                    _j += 5
                    continue
                _fixed_lines.append(f"{_cur}\t\t{_next_content}")
                _j += 3
                continue
            _fixed_lines.append(_broken_lines[_j])
            _j += 1
        lines = _fixed_lines

        i = 0
        # ★ v3.5.1 — 지문↔대사 사이 빈 줄 자동 삽입을 위한 직전 블록 타입 추적
        # 가능한 값: None, "scene", "action", "dialogue", "insert"
        prev_block_type = None
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 빈 줄
            if not stripped:
                i += 1
                continue

            # ★ WRITER_NOTES 마커 블록 스킵 (v3.4 신규)
            # 프롬프트가 "<WRITER_NOTES_BEGIN>...<WRITER_NOTES_END>" 마커로
            # 메타 블록을 감싸도록 지시. 이 마커 안쪽 전체를 차단.
            if "<WRITER_NOTES_BEGIN>" in stripped or "WRITER_NOTES_BEGIN" in stripped:
                i += 1
                while i < len(lines):
                    if ("<WRITER_NOTES_END>" in lines[i] or
                        "WRITER_NOTES_END" in lines[i]):
                        i += 1
                        break
                    i += 1
                continue
            # ★ SPACE_DIVERSITY_CHECK 마커 블록 스킵 (v3.5 신규)
            # 씬 플랜 공간 분산 자가 점검 블록이 시나리오에 유출되는 것 차단
            if ("<SPACE_DIVERSITY_CHECK>" in stripped or
                "SPACE_DIVERSITY_CHECK" in stripped):
                i += 1
                while i < len(lines):
                    if ("</SPACE_DIVERSITY_CHECK>" in lines[i] or
                        "SPACE_DIVERSITY_CHECK" in lines[i] and "</" in lines[i]):
                        i += 1
                        break
                    i += 1
                continue
            # BLOCK 2 헤더 (메타 블록 시작 신호) 감지 — 마커를 잊어도 차단
            if stripped.startswith("[BLOCK 2:") or stripped.startswith("[BLOCK 2 "):
                # 파일 끝까지 스킵
                i = len(lines)
                continue
            if stripped == "[BLOCK 1: 시나리오 본문]" or stripped.startswith("[BLOCK 1:"):
                i += 1
                continue
            if stripped.startswith("━━━"):
                i += 1
                continue

            # 구분선/내부메모 스킵 (--- 또는 ═══ 이후 블록 전체)
            if divider_re.match(stripped):
                i += 1
                # 내부메모 블록 스킵 (다음 씬 헤딩이나 Beat 헤더까지)
                while i < len(lines):
                    memo_line = lines[i].strip()
                    if heading_re.match(memo_line):
                        break
                    if "Beat " in memo_line and "—" in memo_line:
                        break
                    i += 1
                continue

            # 내부 메모 블록 감지 — "**내부 메모**" / "내부 메모:" / "내부 메모" 등
            if "내부 메모" in stripped:
                i += 1
                # 내부 메모 블록 전체 스킵 (다음 씬 헤딩까지)
                while i < len(lines):
                    memo_line = lines[i].strip()
                    if heading_re.match(memo_line):
                        break
                    i += 1
                continue

            # ★ 메타데이터 개별 줄 차단 (v3.4 신규 강화 필터)
            # 장르 장치 이름, 비트 메타 항목, 한글 메타 헤더 등을
            # 광범위하게 매칭하여 본문 유출 차단
            if is_meta_line(stripped):
                i += 1
                continue

            # Beat 헤더 스킵 (프롬프트가 넣는 ACT — Beat 헤더)
            if stripped.startswith("═") or "Beat " in stripped and "—" in stripped:
                i += 1
                continue

            # 씬 헤딩
            m = heading_re.match(stripped)
            if m:
                # S#번호 포함 전체 텍스트 사용
                add_scene_heading(stripped)
                prev_block_type = "scene"  # ★ v3.5.1
                i += 1
                continue

            # 인라인 대사 감지: "캐릭터명\t\t(지시) 대사" 형식 (V.O./O.S. 포함)
            im = inline_dialogue_re.match(stripped)
            if im:
                char_name = im.group(1).strip()
                vo_marker = im.group(2) or ""  # V.O., O.S. 등
                inline_paren = im.group(3) or ""  # (부드럽게) 등
                inline_text = im.group(4).strip()
                # V.O./O.S.를 캐릭터명에 포함
                if vo_marker:
                    char_name = f"{char_name} ({vo_marker})"
                # ★ v3.5.1 — 지문/insert 직후 대사면 빈 줄 1개 삽입
                if prev_block_type in ("action", "insert"):
                    add_blank_line()
                add_dialogue(char_name, inline_paren, inline_text)
                prev_block_type = "dialogue"
                i += 1
                continue

            # 캐릭터명 + 대사 감지 (들여쓰기 형식)
            cm = char_re.match(line)
            if cm:
                char_name = cm.group(1).strip()
                vo_marker = cm.group(2) or ""  # V.O., O.S. 등
                if vo_marker:
                    char_name = f"{char_name} ({vo_marker})"
                parenthetical = ""
                dialogue_lines = []
                # ★ v3.5.1 — 지문/insert 직후 대사면 빈 줄 1개 삽입
                if prev_block_type in ("action", "insert"):
                    add_blank_line()
                i += 1

                # 괄호 지시 확인
                if i < len(lines):
                    pm = paren_re.match(lines[i])
                    if pm:
                        parenthetical = pm.group(1)
                        i += 1

                # 대사 수집
                while i < len(lines):
                    dl = lines[i]
                    ds = dl.strip()
                    if not ds:
                        break
                    if heading_re.match(ds):
                        break
                    if char_re.match(dl):
                        break
                    if inline_dialogue_re.match(ds):
                        break
                    dialogue_lines.append(ds)
                    i += 1

                # 대사 출력 — 한 캐릭터의 연속 대사를 하나로 병합
                # AI가 한 발화 안에서 줄바꿈을 넣으면 여러 줄로 쪼개지는 문제 해결
                if dialogue_lines:
                    # 괄호 지시문 (행동 지시)을 감지해서 분리 처리
                    # 예: "(결제 페이지를 본다)" 같은 줄은 대사와 분리
                    merged_parts = []  # [(type, text)] — type은 "dialogue" 또는 "action"
                    current_dialogue = []
                    for dl in dialogue_lines:
                        dl_stripped = dl.strip()
                        # 줄 전체가 괄호로 시작해서 괄호로 끝나면 행동 지시
                        if (dl_stripped.startswith("(") and dl_stripped.endswith(")")
                            and len(dl_stripped) > 2):
                            # 지금까지 모은 대사를 먼저 합치기
                            if current_dialogue:
                                merged_parts.append(("dialogue", " ".join(current_dialogue)))
                                current_dialogue = []
                            merged_parts.append(("action", dl_stripped))
                        else:
                            current_dialogue.append(dl_stripped)
                    if current_dialogue:
                        merged_parts.append(("dialogue", " ".join(current_dialogue)))

                    # 출력
                    first = True
                    for part_type, part_text in merged_parts:
                        if part_type == "dialogue":
                            if first:
                                add_dialogue(char_name, parenthetical, part_text)
                                parenthetical = ""
                                first = False
                            else:
                                add_dialogue(char_name, "", part_text, continuation=True)
                        else:
                            # 행동 지시는 지문으로 표시 (대사 사이에 끼워 넣기)
                            # ★ v3.5.1 — 대사 직후 행동 지시면 빈 줄 1개 삽입
                            add_blank_line()
                            add_action(part_text)
                            add_blank_line()  # 행동 지시 다음 다시 대사가 오므로 빈 줄
                            # 행동 지시 뒤 대사는 다시 캐릭터명 표시
                            first = True
                    # 처음부터 모두 행동 지시만 있던 경우 fallback
                    if first:
                        add_dialogue(char_name, parenthetical, "")
                else:
                    add_dialogue(char_name, parenthetical, "")
                prev_block_type = "dialogue"  # ★ v3.5.1
                continue

            # 그 외 = 지문
            # ★ v3.5.1 — 대사 직후 지문이면 빈 줄 1개 삽입
            if prev_block_type == "dialogue":
                add_blank_line()
            add_action(stripped)
            prev_block_type = "action"
            i += 1

    # ── 푸터 ──
    doc.add_page_break()
    add_text(f"© 2026 BLUE JEANS PICTURES · Writer Engine {ENGINE_VERSION}",
             size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x8E, 0x8E, 0x99))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════
st.markdown(
    '<div style="text-align:center;padding:1rem 0 0 0">'
    '<div class="header">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>'
    '<div class="brand-title">WRITER ENGINE</div>'
    '<div class="sub">Y O U N G &nbsp; · &nbsp; V I N T A G E &nbsp; · &nbsp; F R E E &nbsp; · &nbsp; I N N O V A T I V E</div>'
    '</div>',
    unsafe_allow_html=True,
)

# ═══════════════════════════════════════════════════════════
# STEP 1 — 자료 입력 (Creator Engine 9개 항목)
# ═══════════════════════════════════════════════════════════
st.markdown(
    '<div class="section-header">📥 STEP 1 · 자료 입력 <span class="en">PASTE FROM CREATOR ENGINE</span></div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="small-meta">Creator Engine 결과를 항목별로 복사해서 붙여넣으세요. 필요한 칸만 채워도 됩니다.</div>',
    unsafe_allow_html=True,
)

# ═══════════════════════════════════════════════════════════
# ★ v3.1 신규 — Creator Engine JSON 자동 로더
# ═══════════════════════════════════════════════════════════
with st.expander("⚡ Creator Engine JSON 업로드 (자동 채우기)", expanded=False):
    st.markdown(
        '<div class="small-meta">Creator Engine에서 내려받은 .json 파일을 업로드하면 '
        '아래 모든 입력칸이 자동으로 채워집니다. BJND·엔딩 페이오프·엔딩 타입도 함께 적용됩니다.</div>',
        unsafe_allow_html=True,
    )
    json_file = st.file_uploader("Creator Engine JSON 파일", type=["json"], key="creator_json_uploader")
    col_u1, col_u2 = st.columns([1, 2])
    with col_u1:
        load_btn = st.button("📂 JSON 적용", use_container_width=True,
                              disabled=(json_file is None))
    with col_u2:
        if json_file is not None:
            st.caption(f"선택됨: `{json_file.name}`")
    
    if load_btn and json_file is not None:
        try:
            raw = json_file.read().decode("utf-8")
            creator_data = json.loads(raw)
            loaded = extract_from_creator_json(creator_data)
            
            # 세션에 주입
            for k, v in loaded.items():
                if k in st.session_state or k in FIELDS:
                    st.session_state[k] = v
            
            # 엔진 버전 표시
            meta = creator_data.get("_meta", {})
            ce_ver = meta.get("engine_version", "?")
            stage = meta.get("stage", "?")
            
            # 엔딩 타입 판정 결과 알림
            et = loaded.get("ending_payoff_type", "")
            if et == "internal_transformation":
                et_msg = "🧠 내적 전환형 (Strategy 전환 — 쿠킹클래스 유형)"
            elif et == "external_choice":
                et_msg = "🎯 외적 선택형 (커플 성사/진실 폭로 등)"
            else:
                et_msg = "⚪ 미판정 (기존 v3.0 장르 엔딩 규칙 사용)"
            
            st.success(
                f"✅ Creator Engine {ce_ver} / {stage} 단계 로드 완료.\n\n"
                f"**프로젝트**: {loaded.get('title', '(무제')})\n"
                f"**엔딩 판정**: {et_msg}\n"
                f"**로드된 필드**: 11칸 + v3.1 신규 3칸"
            )
            st.rerun()
        except json.JSONDecodeError as e:
            st.error(f"JSON 파싱 실패: {e}")
        except Exception as e:
            st.error(f"로드 중 오류: {e}")

# ═══════════════════════════════════════════════════════════
# ★ v3.4.0 신규 — 프로젝트 세션 백업 (저장/불러오기)
# 중간에 멈춰도 처음부터 다시 안 해도 되게.
# ═══════════════════════════════════════════════════════════
with st.expander("💾 프로젝트 세션 백업 (중단 시 복구용)", expanded=False):
    st.markdown(
        '<div class="small-meta">현재 작업 중인 모든 입력칸·씬 플랜·집필된 비트를 JSON으로 저장하거나 불러옵니다. '
        '비트 집필 도중 멈추거나 다음 날 이어서 작업할 때 사용하세요.</div>',
        unsafe_allow_html=True,
    )
    
    col_b1, col_b2 = st.columns(2)
    
    # ── 저장 ──
    with col_b1:
        st.markdown("**📥 백업 저장**")
        _backup_title = st.session_state.get("title", "") or "Untitled"
        _backup_count = len(st.session_state.get("beats_done", {}))
        _backup_bytes = export_session_backup()
        _backup_fname = make_backup_filename(_backup_title, _backup_count)
        st.download_button(
            label=f"💾 JSON 다운로드 ({_backup_count}/15 비트)",
            data=_backup_bytes,
            file_name=_backup_fname,
            mime="application/json",
            use_container_width=True,
            key="backup_download_btn",
        )
        st.caption(f"파일명: `{_backup_fname}`")
    
    # ── 불러오기 ──
    with col_b2:
        st.markdown("**📤 백업 불러오기**")
        backup_file = st.file_uploader(
            "백업 JSON 파일", type=["json"],
            key="backup_uploader",
            label_visibility="collapsed",
        )
        load_backup_btn = st.button(
            "📂 백업 적용 (현재 작업 덮어쓰기)",
            use_container_width=True,
            disabled=(backup_file is None),
            key="backup_load_btn",
        )
        
        if load_backup_btn and backup_file is not None:
            try:
                meta = import_session_backup(backup_file.read())
                
                saved_ver = meta.get("engine_version", "?")
                saved_at = meta.get("saved_at", "?")
                progress = meta.get("beats_progress", "?")
                title = meta.get("title", "(무제)")
                
                # 버전 호환성 안내
                if saved_ver != ENGINE_VERSION:
                    st.warning(
                        f"⚠️ 백업 버전({saved_ver})이 현재 엔진({ENGINE_VERSION})과 다릅니다. "
                        f"복원은 시도되었으나 일부 신규 기능은 반영되지 않을 수 있습니다."
                    )
                
                st.success(
                    f"✅ 백업 복원 완료\n\n"
                    f"**프로젝트**: {title}\n\n"
                    f"**저장 시각**: {saved_at}\n\n"
                    f"**엔진 버전**: {saved_ver}\n\n"
                    f"**진행도**: {progress} 비트"
                )
                st.rerun()
            except json.JSONDecodeError as e:
                st.error(f"JSON 파싱 실패: {e}")
            except Exception as e:
                st.error(f"복원 중 오류: {e}")

col_g1, col_g2 = st.columns(2)
with col_g1:
    # Creator Engine과 동일한 장르 목록 (로맨틱 코미디 포함)
    # _is_comedy/_is_romance가 "로맨틱 코미디"를 자동 감지해 COMEDY+ROMANCE 둘 다 주입
    genre_list = [
        "미지정", "범죄/스릴러", "드라마", "액션", "로맨스", "코미디",
        "로맨틱 코미디", "호러/공포", "SF", "판타지",
        "시대극/사극", "느와르", "미스터리", "전쟁", "뮤지컬", "다큐/논픽션"
    ]
    current_genre = st.session_state.get("genre", "범죄/스릴러")
    if current_genre not in genre_list:
        current_genre = "미지정"
    genre = st.selectbox("장르", genre_list,
                          index=genre_list.index(current_genre))
    st.session_state["genre"] = genre
with col_g2:
    fmt = "영화 (장편)"
    st.session_state["fmt"] = fmt

st.session_state["title"] = st.text_input(
    "프로젝트 제목", value=st.session_state.get("title", ""),
    placeholder="예: 물귀신")

st.session_state["logline"] = st.text_area(
    "Logline", value=st.session_state["logline"],
    height=60, placeholder="Logline Pack (5종 중 택1 또는 전체)")

col_i1, col_i2 = st.columns(2)
with col_i1:
    st.session_state["intent"] = st.text_area(
        "기획의도 (Project Intent)", value=st.session_state["intent"],
        height=100, placeholder="소재 / 장르 / 시장 / Pitch / Theme")
    st.session_state["gns"] = st.text_area(
        "Goal / Need / Strategy", value=st.session_state["gns"],
        height=100, placeholder="Goal / Need / Strategy")
    # ★ v3.6 신규: Creator Engine opening_strategy 전용 필드 (세계관 앞)
    st.session_state["opening_strategy"] = st.text_area(
        "오프닝 전략 (Creator Engine — OPENING STRATEGY)",
        value=st.session_state.get("opening_strategy", ""),
        height=140,
        placeholder="Creator Engine의 오프닝 전략 6개 필드를 그대로 붙여넣으세요.\n"
                    "[오프닝 타입] / [오프닝 의도] / [도파민 포인트] / "
                    "[1막 연결 방식] / [훅 라인/이미지] / [장르 DNA 체크]\n"
                    "비워두면 장르 일반 가이드로 폴백합니다.",
        help="Creator Engine Core Build 결과물의 OPENING STRATEGY 섹션을 "
             "그대로 복사해서 붙여넣으세요. Beat 1 집필 시 이 구체적 전략이 "
             "장르 일반 가이드보다 우선 적용됩니다."
    )
    st.session_state["world"] = st.text_area(
        "세계관 (World Build)", value=st.session_state["world"],
        height=100, placeholder="시간 / 공간 / 규칙 / 금기 / 권력구조")
with col_i2:
    st.session_state["characters"] = st.text_area(
        "캐릭터 + 바이블 ← 가장 중요", value=st.session_state["characters"],
        height=300, placeholder="캐릭터(4인) + 바이블(백스토리/말투규칙/대사샘플/관계태도/변화궤적)")

st.session_state["structure"] = st.text_area(
    "구조 (Synopsis + Storyline + Beat Sheet)", value=st.session_state["structure"],
    height=120, placeholder="Synopsis 1P / Storyline 8시퀀스 / 15-Beat Sheet / 3막 진단")

st.session_state["scene_design"] = st.text_area(
    "장면 설계 (Scene Design)", value=st.session_state["scene_design"],
    height=120, placeholder="핵심 장면 15~18개 + Scene Map")

st.session_state["treatment"] = st.text_area(
    "트리트먼트 (Treatment)", value=st.session_state["treatment"],
    height=160, placeholder="16비트 줄글 (1막/2막/3막)")

st.session_state["tone"] = st.text_area(
    "톤 문서 (Tone Document)", value=st.session_state["tone"],
    height=80, placeholder="비주얼/페이싱/대사규칙/모티프/사운드/금기/Writer지시")

# ═══════════════════════════════════════════════════════════
# ★ v3.1 신규 — BJND · Ending Payoff · Ending Type
# Creator JSON 업로드 시 자동 채워짐. 수동 편집도 가능.
# ═══════════════════════════════════════════════════════════
with st.expander("🧬 v3.1 BJND · 엔딩 설계 (Creator Engine v2.3+ 연동)", expanded=False):
    st.markdown(
        '<div class="small-meta">Creator Engine JSON을 업로드하면 자동 채워집니다. '
        '수동 편집도 가능합니다. 이 세 항목이 비어 있으면 기존 v3.0 방식(장르 엔딩 규칙)으로 집필됩니다.</div>',
        unsafe_allow_html=True,
    )
    
    st.session_state["bjnd_data"] = st.text_area(
        "BJND 설계 (Loss/Lack · Goal · Need · Strategy · Cost · Ending Payoff)",
        value=st.session_state.get("bjnd_data", ""),
        height=150,
        placeholder=(
            "[desire_origin] lack\n"
            "[goal] 독립적 삶 증명\n"
            "[need] 간보지 않는 친밀함\n"
            "[strategy] 두 남자를 재료처럼 비교·분석\n"
            "[cost_1] 관계의 피상성\n"
            "[ending_payoff] ..."
        ),
        help="매 비트에서 Strategy가 행동으로 드러나도록 강제하고, 비트 구간별 Cost 단계를 강제합니다."
    )
    
    col_e1, col_e2 = st.columns([3, 2])
    with col_e1:
        st.session_state["ending_payoff"] = st.text_area(
            "Ending Payoff (엔딩 페이오프 텍스트)",
            value=st.session_state.get("ending_payoff", ""),
            height=80,
            placeholder="예: 유진이 분석을 멈추고 세웅의 이름을 먼저 부르는 순간, Need(친밀함)가 처음으로 채워진다.",
            help="Beat 15~16에서 이 페이오프가 구체적 행동/이미지로 구현되도록 강제합니다."
        )
    with col_e2:
        et_options = ["(미지정 — 장르 폴백)", "internal_transformation", "external_choice"]
        et_labels = {
            "(미지정 — 장르 폴백)": "",
            "internal_transformation": "internal_transformation",
            "external_choice": "external_choice",
        }
        current_et = st.session_state.get("ending_payoff_type", "")
        # 역매핑
        display_et = "(미지정 — 장르 폴백)"
        for label, val in et_labels.items():
            if val == current_et and val:
                display_et = label
                break
        
        chosen = st.selectbox(
            "엔딩 타입 (Ending Type)",
            options=et_options,
            index=et_options.index(display_et) if display_et in et_options else 0,
            help=(
                "internal_transformation: 자기 발견·Strategy 전환형 엔딩 "
                "(쿠킹클래스 유형 — 한 명 선택 강제 해제)\n\n"
                "external_choice: 명확한 외적 선택·대결·고백형 엔딩\n\n"
                "미지정: 기존 v3.0 장르 엔딩 규칙 사용"
            )
        )
        st.session_state["ending_payoff_type"] = et_labels[chosen]
    
    # 엔딩 판정 결과 실시간 표시
    et_now = st.session_state.get("ending_payoff_type", "")
    if et_now == "internal_transformation":
        st.info("🧠 **내적 전환형 엔딩** — Beat 15~16에서 '한 명 선택' 엔딩 금지, Strategy 전환 씬 강제")
    elif et_now == "external_choice":
        st.info("🎯 **외적 선택형 엔딩** — 장르 약속에 따른 명확한 외적 행동 엔딩")
    else:
        st.caption("⚪ 미지정 상태 — 기존 v3.0 장르 기반 엔딩 규칙이 Beat 15에 적용됩니다.")

# ── 실화 배경 기반 작품 체크 ──
st.session_state["fact_based"] = st.checkbox(
    "실화 배경 기반 작품 (실명 비사용 + 사실성 확보 규칙 적용)",
    value=st.session_state.get("fact_based", False),
    help="체크하면 집필 시 실명·특정 가능 디테일 회피 규칙이 적용되고, "
         "DOCX에 각색 고지 자막이 자동 삽입됩니다. "
         "실제 지명·시대 사건·공적 직함은 사용 가능합니다."
)

# ── 역사영화 체크 + 유형 선택 ──
_hist_col1, _hist_col2 = st.columns([2, 3])
with _hist_col1:
    st.session_state["historical"] = st.checkbox(
        "역사영화 (시대 고증 + 유형별 규칙 적용)",
        value=st.session_state.get("historical", False),
        help="체크하면 시대감 구체화·시대 언어 균형·공간 설계 등 역사영화 공통 규칙이 적용되고, "
             "선택한 유형(정통/팩션/퓨전)에 따라 세부 집필 철학이 분기됩니다."
    )
with _hist_col2:
    if st.session_state.get("historical", False):
        st.session_state["historical_type"] = st.selectbox(
            "역사영화 유형",
            options=["정통역사영화", "팩션역사영화", "퓨전역사영화"],
            index=["정통역사영화", "팩션역사영화", "퓨전역사영화"].index(
                st.session_state.get("historical_type", "팩션")
                if st.session_state.get("historical_type", "팩션") in ["정통역사영화", "팩션역사영화", "퓨전역사영화"]
                else ("정통역사영화" if "정통" in st.session_state.get("historical_type", "팩션")
                      else "팩션역사영화" if "팩션" in st.session_state.get("historical_type", "팩션")
                      else "퓨전역사영화")
            ),
            help="정통: 실재 사건·인물 중심, 사료 존중, 감정 과잉 회피 (남한산성·1987)  |  "
                 "팩션: 실재 시대+가공 드라마, 재미 코드 자유 (왕의 남자·암살·밀정)  |  "
                 "퓨전: 시대 차용+자유 서사, 장르 재미 우선 (조선명탐정·전우치)"
        )

# API 상태
if get_client():
    st.success(f"API 준비 완료 — 집필: {ANTHROPIC_MODEL_WRITE} · 구조: {ANTHROPIC_MODEL_PLAN}")
else:
    st.warning("ANTHROPIC_API_KEY가 설정되지 않았습니다.")

has_material = any(st.session_state[f].strip() for f in FIELDS)

# ═══════════════════════════════════════════════════════════
# SCENE PLAN — 3막 분할 (100씬 / 100분)
# ═══════════════════════════════════════════════════════════
st.markdown(
    '<div class="section-header">🗺️ 씬 플랜 <span class="en">SCENE PLAN · 100 SCENES / 3-ACT SPLIT</span></div>',
    unsafe_allow_html=True,
)

if has_material:
    st.markdown(
        '<div class="small-meta">'
        '100씬/100분 기준. 1막 → 2막 → 3막 순서로 생성합니다.'
        '</div>',
        unsafe_allow_html=True,
    )

    plan_kw = dict(
        genre=genre, fmt=fmt,
        logline=st.session_state["logline"],
        intent=st.session_state["intent"],
        gns=st.session_state["gns"],
        characters=st.session_state["characters"],
        opening_strategy=st.session_state.get("opening_strategy", ""),  # ★ v3.6 신규
        world=st.session_state["world"],
        structure=st.session_state["structure"],
        scene_design=st.session_state["scene_design"],
        treatment=st.session_state["treatment"],
        tone=st.session_state["tone"],
        fact_based=st.session_state.get("fact_based", False),
        historical=st.session_state.get("historical", False),
        historical_type=st.session_state.get("historical_type", "팩션"),
    )

    col_p1, col_p2, col_p3 = st.columns(3)

    with col_p1:
        done1 = bool(st.session_state["plan_1막"])
        btn1 = st.button(
            f"{'✅ ' if done1 else ''}1막 플랜 (Beat 1~5)",
            type="primary" if not done1 else "secondary",
            use_container_width=True,
        )
    with col_p2:
        done2 = bool(st.session_state["plan_2막"])
        btn2 = st.button(
            f"{'✅ ' if done2 else ''}2막 플랜 (Beat 6~11)",
            type="primary" if done1 and not done2 else "secondary",
            use_container_width=True,
            disabled=not done1,
        )
    with col_p3:
        done3 = bool(st.session_state["plan_3막"])
        btn3 = st.button(
            f"{'✅ ' if done3 else ''}3막 플랜 (Beat 12~15)",
            type="primary" if done2 and not done3 else "secondary",
            use_container_width=True,
            disabled=not done2,
        )

    if btn1:
        prompt = build_scene_plan_prompt(act="1막", **plan_kw, previous_plan="")
        st.markdown('<div class="act-tag">1막 씬 플랜 생성 중…</div>', unsafe_allow_html=True)
        result = st.write_stream(stream_ai(prompt, model=ANTHROPIC_MODEL_PLAN))
        st.session_state["plan_1막"] = result
        st.session_state["plan_2막"] = ""
        st.session_state["plan_3막"] = ""
        st.session_state["beats_done"] = {}
        st.session_state["current_beat"] = 1
        st.rerun()

    if btn2:
        prompt = build_scene_plan_prompt(act="2막", **plan_kw, previous_plan=st.session_state["plan_1막"])
        st.markdown('<div class="act-tag">2막 씬 플랜 생성 중…</div>', unsafe_allow_html=True)
        result = st.write_stream(stream_ai(prompt, model=ANTHROPIC_MODEL_PLAN))
        st.session_state["plan_2막"] = result
        st.session_state["plan_3막"] = ""
        st.rerun()

    if btn3:
        prev = st.session_state["plan_1막"] + "\n\n" + st.session_state["plan_2막"]
        prompt = build_scene_plan_prompt(act="3막", **plan_kw, previous_plan=prev)
        st.markdown('<div class="act-tag">3막 씬 플랜 생성 중…</div>', unsafe_allow_html=True)
        result = st.write_stream(stream_ai(prompt, model=ANTHROPIC_MODEL_PLAN))
        st.session_state["plan_3막"] = result
        st.rerun()

    # 플랜 표시
    for act in ["1막", "2막", "3막"]:
        p = st.session_state.get(f"plan_{act}", "")
        if p:
            with st.expander(f"{act} 씬 플랜 ✅", expanded=False):
                st.text(p)

    if plan_ready():
        st.markdown(
            '<div class="callout"><div class="cl">PLAN COMPLETE</div>'
            '3막 씬 플랜 완성. 아래에서 비트별 집필을 시작하세요.</div>',
            unsafe_allow_html=True,
        )
        plan_all = full_plan()
        st.download_button(
            label="씬 플랜 TXT 저장",
            data=plan_all,
            file_name=f"scene_plan_{genre}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
            mime="text/plain",
            use_container_width=True,
        )

        # ── 핵심 요소 추출 ──
        st.markdown(
            '<div class="section-header">🔍 핵심 요소 추출 <span class="en">STORY ELEMENTS</span></div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            '<div class="small-meta">'
            '기획 자료에서 맥거핀·캐릭터 비밀·핵심 장소·모티프·클라이맥스 설계를 추출합니다. '
            '추출 결과는 매 비트 집필 시 자동으로 주입되어 누락을 방지합니다.'
            '</div>',
            unsafe_allow_html=True,
        )

        elements_done = bool(st.session_state.get("story_elements", ""))

        if st.button(
            f"{'✅ ' if elements_done else ''}핵심 요소 추출",
            type="primary" if not elements_done else "secondary",
            use_container_width=True,
        ):
            extract_prompt = build_extract_elements_prompt(
                genre=genre,
                logline=st.session_state["logline"],
                characters=st.session_state["characters"],
                structure=st.session_state["structure"],
                scene_design=st.session_state["scene_design"],
                treatment=st.session_state["treatment"],
                tone=st.session_state["tone"],
                world=st.session_state["world"],
            )
            st.markdown('<div class="beat-tag">핵심 요소 추출 중…</div>', unsafe_allow_html=True)
            result = st.write_stream(stream_ai(extract_prompt, model=ANTHROPIC_MODEL_PLAN))
            st.session_state["story_elements"] = result
            st.rerun()

        if elements_done:
            with st.expander("핵심 요소 보기 ✅", expanded=False):
                st.text(st.session_state["story_elements"])
else:
    st.markdown(
        '<div class="callout"><div class="cl">WAITING</div>'
        '위에 기획 자료를 붙여넣으면 시작할 수 있습니다.</div>',
        unsafe_allow_html=True,
    )

# ═══════════════════════════════════════════════════════════
# STEP 2 — 비트별 집필
# ═══════════════════════════════════════════════════════════
if plan_ready():
    st.markdown(
        '<div class="section-header">✍️ STEP 2 · 비트별 집필 <span class="en">WRITE BY BEAT</span></div>',
        unsafe_allow_html=True,
    )

    if not st.session_state.get("story_elements"):
        st.warning("⚠️ 핵심 요소 추출을 먼저 실행하세요. 맥거핀·비밀·모티프가 누락될 수 있습니다.")

    cur = st.session_state["current_beat"]
    done = st.session_state["beats_done"]
    combined_plan = full_plan()

    # 완료 비트 표시
    for b_no in sorted(done.keys()):
        b_info = BEATS_15[b_no - 1]
        with st.expander(
            f"{b_info['act']} — Beat {b_no}. {b_info['name']} ✅",
            expanded=(b_no == max(done.keys())),
        ):
            st.text(done[b_no])

    # 현재 비트 정보
    if cur <= 15:
        b_info = BEATS_15[cur - 1]
        st.markdown(
            f'<div class="beat-tag">Beat {cur} / 15</div> '
            f'<span style="font-weight:700">{b_info["name"]}</span> '
            f'<span style="color:var(--dim)">({b_info["act"]})</span>',
            unsafe_allow_html=True,
        )

    col_b1, col_b2 = st.columns(2)
    with col_b1:
        write_btn = st.button(
            f"Beat {cur} 집필" if cur <= 15 else "전체 완료 ✅",
            type="primary", use_container_width=True,
            disabled=(cur > 15),
        )
    with col_b2:
        rewrite_btn = st.button(
            "마지막 비트 다시",
            use_container_width=True,
            disabled=(len(done) == 0),
        )

    rewrite_note = ""
    if rewrite_btn:
        rewrite_note = st.text_input(
            "수정 지시 (비워두면 전체 강화)",
            placeholder="예: 대사를 더 차갑게 / 긴장감 올려줘 / Hook 강화",
        )

    # 집필
    if write_btn and cur <= 15:
        prev_text = done[cur - 1] if (cur > 1 and (cur - 1) in done) else ""
        prompt = build_write_beat_prompt(
            genre=genre, beat_number=cur,
            scene_plan=combined_plan,
            characters=st.session_state["characters"],
            treatment=st.session_state["treatment"],
            tone=st.session_state["tone"],
            previous_scene_text=prev_text,
            logline=st.session_state["logline"],
            world=st.session_state["world"],
            story_elements=st.session_state.get("story_elements", ""),
            opening_strategy=st.session_state.get("opening_strategy", ""),  # v3.6
            bjnd_data=st.session_state.get("bjnd_data", ""),                  # ★ v3.1
            ending_payoff=st.session_state.get("ending_payoff", ""),          # ★ v3.1
            ending_payoff_type=st.session_state.get("ending_payoff_type", ""),# ★ v3.1
            fact_based=st.session_state.get("fact_based", False),
            historical=st.session_state.get("historical", False),
            historical_type=st.session_state.get("historical_type", "팩션"),
        )
        st.markdown(f'<div class="beat-tag">Beat {cur} 집필 중…</div>', unsafe_allow_html=True)
        result = st.write_stream(stream_ai(prompt, tokens=16000))
        st.session_state["beats_done"][cur] = result
        st.session_state["current_beat"] = cur + 1
        st.rerun()

    # 다시 쓰기
    if rewrite_btn and done:
        last_beat = max(done.keys())
        prompt = build_rewrite_prompt(
            genre=genre, beat_number=last_beat,
            current_text=done[last_beat],
            characters=st.session_state["characters"],
            instruction=rewrite_note,
        )
        st.markdown(f'<div class="beat-tag">Beat {last_beat} 다시 쓰는 중…</div>', unsafe_allow_html=True)
        result = st.write_stream(stream_ai(prompt, tokens=16000))
        st.session_state["beats_done"][last_beat] = result
        st.rerun()

# ═══════════════════════════════════════════════════════════
# DOWNLOAD — TXT + DOCX (수시 저장)
# ═══════════════════════════════════════════════════════════
if st.session_state.get("beats_done"):
    st.markdown(
        '<div class="section-header">📄 다운로드 <span class="en">EXPORT · SAVE ANYTIME</span></div>',
        unsafe_allow_html=True,
    )

    done_count = len(st.session_state["beats_done"])
    st.markdown(
        f'<div class="callout"><div class="cl">DATA</div>'
        f'{done_count}/15 비트 완료. 새로고침하면 데이터가 사라집니다. 수시로 저장하세요.</div>',
        unsafe_allow_html=True,
    )

    # TXT
    parts = []
    for b_no in sorted(st.session_state["beats_done"].keys()):
        b_info = BEATS_15[b_no - 1]
        # ★ v3.5.1 — 지문↔대사 빈 줄 후처리 적용
        beat_text = _normalize_screenplay_blank_lines(
            st.session_state['beats_done'][b_no]
        )
        parts.append(
            f"{'='*60}\n{b_info['act']} — Beat {b_no}. {b_info['name']}\n{'='*60}\n\n"
            f"{beat_text}"
        )
    all_text = "\n\n\n".join(parts)

    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button(
            label=f"TXT 저장 ({done_count}/15)",
            data=all_text,
            file_name=_build_download_filename(
                st.session_state.get("title", ""), genre, "txt"
            ),
            mime="text/plain",
            use_container_width=True,
        )
    with col_dl2:
        try:
            docx_bytes = make_docx_bytes(genre, st.session_state["beats_done"],
                                         title=st.session_state.get("title", ""),
                                         fact_based=st.session_state.get("fact_based", False),
                                         historical=st.session_state.get("historical", False),
                                         historical_type=st.session_state.get("historical_type", "팩션"))
            st.download_button(
                label=f"DOCX 저장 ({done_count}/15)",
                data=docx_bytes,
                file_name=_build_download_filename(
                    st.session_state.get("title", ""), genre, "docx"
                ),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.caption("DOCX: python-docx 미설치 — pip install python-docx")

# ═══════════════════════════════════════════════════════════
# RESET
# ═══════════════════════════════════════════════════════════
st.markdown("---")
col_r1, col_r2 = st.columns([3, 1])
with col_r2:
    if st.button("전체 초기화", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

st.caption("© 2026 BLUE JEANS PICTURES · Writer Engine v3.0")
